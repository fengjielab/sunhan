#!/usr/bin/env python3
"""
Format .docx paper to match 《制造业自动化》 journal template.
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import copy
import os

# ── Paths ──────────────────────────────────────────────
SRC = r'F:\前途文件\my_test\data\中文核心投稿稿_视觉语义驱动多参数阻抗辅助遥操作方法_格式整理前备份.docx'
DST = r'F:\前途文件\my_test\data\中文核心投稿稿_视觉语义驱动多参数阻抗辅助遥操作方法_格式整理后.docx'

# ── Font size constants (in Pt) ────────────────────────
# 二号=22pt, 四号=14pt, 小四号=12pt, 五号=10.5pt
SIZE_ER_HAO = Pt(22)      # 二号
SIZE_SI_HAO = Pt(14)      # 四号
SIZE_XIAO_SI = Pt(12)     # 小四号
SIZE_WU = Pt(10.5)        # 五号

# ── Placeholder texts ──────────────────────────────────
AUTHORS_TEXT = '作者1¹, 作者2¹˂*'
AFFILIATIONS_TEXT = '(1. XX大学 XX学院，XX市 邮编；2. XX公司，XX市 邮编) [待补充]'
CLC_TEXT = 'TP242 [待确认]'
EN_TITLE_TEXT = 'Vision-Semantic-Driven Multi-Parameter Impedance-Assisted Teleoperation Method [待确认]'
EN_AUTHORS_TEXT = 'AUTHOR One¹, AUTHOR Two¹˂* [待确认]'
EN_ABSTRACT_TEXT = '[English abstract to be provided — should correspond to the Chinese abstract. Please replace this placeholder with the translated abstract.]'
EN_KEYWORDS_TEXT = 'teleoperation; visual semantics; impedance control; force feedback; human-robot interaction; Franka Panda [待确认]'
AUTHOR_BIO_TEXT = '作者简介：[姓名]（[出生年]-），[性别]，[民族]，[籍贯]人，[职称]，[学位]，研究方向为[研究方向]。'
FUNDING_TEXT = '基金项目：[基金类型 基金编号（基金名称）] [待补充]'


# ═══════════════════════════════════════════════════════
# Helper functions
# ═══════════════════════════════════════════════════════

def set_run_fonts(run, western='Times New Roman', east_asian='宋体',
                  size=None, bold=None):
    """Set both Western and East-Asian fonts, size, bold on a run."""
    rPr = run._element.get_or_add_rPr()

    # Remove existing rFonts if present, or create new
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    # Clear theme font references
    for attr in ['w:asciiTheme', 'w:eastAsiaTheme', 'w:hAnsiTheme', 'w:cstheme']:
        a = rFonts.get(qn(attr))
        if a is not None:
            del rFonts.attrib[qn(attr)]

    rFonts.set(qn('w:ascii'), western)
    rFonts.set(qn('w:hAnsi'), western)
    rFonts.set(qn('w:eastAsia'), east_asian)

    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold


def set_paragraph_format(paragraph, line_spacing=None, space_before=None,
                         space_after=None, alignment=None, first_line_indent=None):
    """Set paragraph-level formatting."""
    pf = paragraph.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if alignment is not None:
        paragraph.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def format_all_runs(paragraph, western='Times New Roman', east_asian='宋体',
                    size=None, bold=None):
    """Apply font settings to every run in a paragraph."""
    for run in paragraph.runs:
        set_run_fonts(run, western=western, east_asian=east_asian,
                      size=size, bold=bold)


def clear_run_formatting(run):
    """Remove explicit formatting from a run so style defaults apply."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        # Remove font name references
        for tag in [qn('w:rFonts'), qn('w:sz'), qn('w:b'), qn('w:i'),
                     qn('w:color'), qn('w:u'), qn('w:vertAlign')]:
            el = rPr.find(tag)
            if el is not None:
                rPr.remove(el)


def insert_paragraph_after(ref_paragraph, new_paragraph):
    """Insert a new paragraph element right after ref_paragraph."""
    ref_paragraph._element.addnext(new_paragraph._element)


def create_new_paragraph(doc, text, western='Times New Roman', east_asian='宋体',
                         size=None, bold=None, alignment=None,
                         line_spacing=None, space_before=None, space_after=None,
                         first_line_indent=None):
    """Create a fully formatted paragraph and return it."""
    p = doc.add_paragraph()
    # Remove from end of document — caller will re-insert at correct position
    doc.element.body.remove(p._element)

    run = p.add_run(text)
    set_run_fonts(run, western=western, east_asian=east_asian, size=size, bold=bold)
    set_paragraph_format(p, line_spacing=line_spacing, space_before=space_before,
                         space_after=space_after, alignment=alignment,
                         first_line_indent=first_line_indent)
    return p


def is_heading_1(text):
    """Check if text is a level-1 heading (e.g. '1 引言')"""
    return bool(re.match(r'^\d+\s+\S+', text))


def is_heading_2(text):
    """Check if text is a level-2 heading (e.g. '3.1 物体操作属性划分')"""
    return bool(re.match(r'^\d+\.\d+\s+\S+', text))


def is_heading_3(text):
    """Check if text is a level-3 heading (e.g. '1.1.1 xxx')"""
    return bool(re.match(r'^\d+\.\d+\.\d+\s+\S+', text))


def is_table_caption(text):
    """Check if text is a table caption (e.g. '表 1 xxx') — short title, not body text."""
    if not re.match(r'^表\s*\d+', text):
        return False
    # Body text like "表 3 给出了..." or "由表 3 可见..." is NOT a caption
    body_verbs = ['给出', '可见', '可看', '显示', '表明', '列出', '汇总', '可看出', '可以看出']
    for verb in body_verbs:
        if verb in text:
            return False
    # Captions are typically shorter (just table number + title, no period in title)
    return True


def is_figure_caption(text):
    """Check if text is a figure caption (e.g. '【图 1 ...' or '图 1 ...')"""
    return bool(re.match(r'^[【]?图\s*\d+', text))


def is_reference_item(text):
    """Check if text is a single reference item (e.g. '[1] Hogan N...')"""
    return bool(re.match(r'^\[\d+\]', text))


# ═══════════════════════════════════════════════════════
# Main formatting logic
# ═══════════════════════════════════════════════════════

def main():
    print("Loading document...")
    doc = Document(SRC)

    # ── 1. Page setup ──────────────────────────────────
    print("Setting page setup...")
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Ensure single column (删掉可能存在的分栏设置)
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is not None:
        cols.set(qn('w:num'), '1')

    # ── 2. Collect paragraph list ──────────────────────
    # We'll work with a copy of the paragraph list since we'll be modifying
    paragraphs = list(doc.paragraphs)

    # ── 3. Process each paragraph ──────────────────────
    print(f"Processing {len(paragraphs)} paragraphs...")

    refs_heading_idx = None
    refs_content_idx = None
    keywords_idx = None

    for idx, p in enumerate(paragraphs):
        text = p.text.strip()
        style = p.style.name

        # ── 3a. Title (段落0: Heading 1) ────────────────
        if idx == 0:
            print(f"  P{idx}: Chinese Title")
            format_all_runs(p, east_asian='黑体', size=SIZE_ER_HAO, bold=True)
            set_paragraph_format(p, line_spacing=1.0,
                                 space_before=SIZE_ER_HAO,
                                 space_after=SIZE_ER_HAO,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # ── 3b. Abstract heading ────────────────────────
        elif text == '摘要':
            print(f"  P{idx}: Abstract heading")
            format_all_runs(p, east_asian='黑体', size=SIZE_XIAO_SI, bold=True)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # ── 3c. Abstract content ───────────────────────
        elif idx > 0 and paragraphs[idx-1].text.strip() == '摘要':
            print(f"  P{idx}: Abstract content")
            format_all_runs(p, east_asian='宋体', size=SIZE_WU)
            set_paragraph_format(p, line_spacing=1.5,
                                 alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # ── 3d. Keywords ───────────────────────────────
        elif text.startswith('关键词') or text.startswith('關键词') or text.startswith('关键词'):
            print(f"  P{idx}: Keywords")
            keywords_idx = idx
            # Split into heading and content properly
            full_text = p.text
            # Find the colon separator
            colon_pos = -1
            for sep in ['：', ':']:
                cp = full_text.find(sep)
                if cp != -1:
                    colon_pos = cp
                    break
            if colon_pos != -1:
                heading_text = full_text[:colon_pos + 1]  # e.g. "关键词："
                content_text = full_text[colon_pos + 1:]   # e.g. "遥操作；视觉语义..."
                # Clear all existing runs
                for run in p.runs:
                    run.text = ''
                # Set first run to heading
                if p.runs:
                    p.runs[0].text = heading_text
                    set_run_fonts(p.runs[0], east_asian='黑体', size=SIZE_XIAO_SI, bold=True)
                    # Add remaining content to second run (or create if needed)
                    if len(p.runs) >= 2:
                        p.runs[1].text = content_text
                        set_run_fonts(p.runs[1], east_asian='宋体', size=SIZE_WU)
                    else:
                        from docx.oxml import OxmlElement as OxE
                        new_r = OxE('w:r')
                        p.runs[0]._element.addnext(new_r)
                        from docx.text.run import Run
                        content_run = Run(new_r, p)
                        content_run.text = content_text
                        set_run_fonts(content_run, east_asian='宋体', size=SIZE_WU)
                # Clear extra runs
                for run in p.runs[2:]:
                    run.text = ''
            else:
                format_all_runs(p, east_asian='黑体', size=SIZE_XIAO_SI, bold=True)
            set_paragraph_format(p, line_spacing=1.5)

        # ── 3e. Section level-1 heading ─────────────────
        elif (style == 'Heading 2' or style == '标题 2') and is_heading_1(text):
            print(f"  P{idx}: L1 heading - {text[:40]}")
            format_all_runs(p, east_asian='黑体', size=SIZE_SI_HAO, bold=True)
            set_paragraph_format(p, space_before=Pt(6), space_after=Pt(3))

        # ── 3f. Section level-2 heading ─────────────────
        elif (style == 'Heading 3' or style == '标题 3') and is_heading_2(text):
            print(f"  P{idx}: L2 heading - {text[:40]}")
            format_all_runs(p, east_asian='黑体', size=SIZE_XIAO_SI, bold=True)
            set_paragraph_format(p, space_before=Pt(3), space_after=Pt(2))

        # ── 3g. References heading ──────────────────────
        elif text == '参考文献' or text == '參考文献' or text == '参考文献':
            print(f"  P{idx}: References heading")
            refs_heading_idx = idx
            format_all_runs(p, east_asian='黑体', size=SIZE_XIAO_SI, bold=True)
            set_paragraph_format(p, space_before=Pt(12))

        # ── 3h. References content ──────────────────────
        elif refs_heading_idx is not None and idx == refs_heading_idx + 1:
            print(f"  P{idx}: References content (will split later)")
            refs_content_idx = idx
            # Don't format yet — will be split

        # ── 3i. Figure captions ─────────────────────────
        elif is_figure_caption(text):
            print(f"  P{idx}: Figure caption")
            # Clean up 【图 N ...】 → 图 N  ...
            clean_text = re.sub(r'^[【]\s*(图\s*\d+)', r'\1', text)
            clean_text = re.sub(r'[】]$', '', clean_text)
            clean_text = re.sub(r'。建议.*$', '。', clean_text)
            # Replace all text in runs
            if p.runs:
                for run in p.runs:
                    run.text = ''
                p.runs[0].text = clean_text
            format_all_runs(p, east_asian='宋体', size=SIZE_WU)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_before=Pt(6), space_after=Pt(6))

        # ── 3j. Table captions ──────────────────────────
        elif is_table_caption(text):
            print(f"  P{idx}: Table caption")
            format_all_runs(p, east_asian='宋体', size=SIZE_WU, bold=True)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_before=Pt(6), space_after=Pt(3))

        # ── 3k. Body text / everything else ─────────────
        else:
            print(f"  P{idx}: Body text ({text[:50]}...)")
            format_all_runs(p, east_asian='宋体', size=SIZE_XIAO_SI)
            set_paragraph_format(p, line_spacing=1.5,
                                 alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                 first_line_indent=SIZE_XIAO_SI * 2)  # 2-char indent

    # ── 4. Insert missing sections after keywords ──────
    print("Inserting English sections and metadata...")

    insert_paragraphs = []

    # 中图分类号
    insert_paragraphs.append(
        ('中图分类号', CLC_TEXT, '黑体', SIZE_XIAO_SI, '宋体', SIZE_WU))
    # Empty line
    insert_paragraphs.append(('spacer', None, None, None, None, None))
    # English title
    insert_paragraphs.append(
        ('title', EN_TITLE_TEXT, 'Times New Roman', SIZE_ER_HAO, None, None))
    # English authors
    insert_paragraphs.append(
        ('authors', EN_AUTHORS_TEXT, 'Times New Roman', SIZE_WU, None, None))
    # English abstract heading + content
    insert_paragraphs.append(
        ('abstract_heading', 'Abstract: ' + EN_ABSTRACT_TEXT, 'Times New Roman', SIZE_WU, None, None))
    # English keywords
    insert_paragraphs.append(
        ('keywords', 'Key words: ' + EN_KEYWORDS_TEXT, 'Times New Roman', SIZE_WU, None, None))

    # Find insertion point (after keywords paragraph)
    if keywords_idx is not None:
        insert_after = paragraphs[keywords_idx]
    else:
        insert_after = paragraphs[3]  # fallback

    new_paras = []
    for item_type, content, font_w, size_w, font_e, size_e in insert_paragraphs:
        if item_type == 'spacer':
            p = doc.add_paragraph()
            doc.element.body.remove(p._element)
            new_paras.append(p)
            continue

        p = doc.add_paragraph()
        doc.element.body.remove(p._element)

        if item_type == 'title':
            run = p.add_run(content)
            set_run_fonts(run, western='Times New Roman', east_asian='Times New Roman',
                          size=SIZE_ER_HAO, bold=True)
            set_paragraph_format(p, line_spacing=1.0, space_before=SIZE_ER_HAO,
                                 space_after=SIZE_ER_HAO,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)
        elif item_type == 'authors':
            run = p.add_run(content)
            set_run_fonts(run, western='Times New Roman', east_asian='Times New Roman',
                          size=SIZE_WU)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        elif item_type == 'abstract_heading':
            run_heading = p.add_run('Abstract: ')
            set_run_fonts(run_heading, western='Times New Roman',
                          east_asian='Times New Roman', size=SIZE_XIAO_SI, bold=True)
            run_content = p.add_run(EN_ABSTRACT_TEXT)
            set_run_fonts(run_content, western='Times New Roman',
                          east_asian='Times New Roman', size=SIZE_WU)
            set_paragraph_format(p, line_spacing=1.5)
        elif item_type == 'keywords':
            run_heading = p.add_run('Key words: ')
            set_run_fonts(run_heading, western='Times New Roman',
                          east_asian='Times New Roman', size=SIZE_XIAO_SI, bold=True)
            run_content = p.add_run(EN_KEYWORDS_TEXT)
            set_run_fonts(run_content, western='Times New Roman',
                          east_asian='Times New Roman', size=SIZE_WU)
            set_paragraph_format(p, space_after=SIZE_WU)
        elif item_type == '中图分类号':
            run_heading = p.add_run('中图分类号：')
            set_run_fonts(run_heading, east_asian='黑体', size=SIZE_XIAO_SI, bold=True)
            run_content = p.add_run(CLC_TEXT)
            set_run_fonts(run_content, east_asian='宋体', size=SIZE_WU)

        new_paras.append(p)

    # Insert all new paragraphs in reverse order after insert_after
    for p in reversed(new_paras):
        insert_paragraph_after(insert_after, p)

    # Also insert Chinese authors/affiliations after title
    print("Inserting Chinese authors and affiliations...")
    author_paras = []

    # Authors
    p_auth = doc.add_paragraph()
    doc.element.body.remove(p_auth._element)
    run = p_auth.add_run(AUTHORS_TEXT)
    set_run_fonts(run, east_asian='宋体', size=SIZE_WU)
    set_paragraph_format(p_auth, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    author_paras.append(p_auth)

    # Affiliations
    p_aff = doc.add_paragraph()
    doc.element.body.remove(p_aff._element)
    run = p_aff.add_run(AFFILIATIONS_TEXT)
    set_run_fonts(run, east_asian='宋体', size=SIZE_WU)
    set_paragraph_format(p_aff, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    author_paras.append(p_aff)

    # Insert after title (paragraph 0)
    for p in reversed(author_paras):
        insert_paragraph_after(paragraphs[0], p)

    # ── 5. Split references ────────────────────────────
    if refs_content_idx is not None:
        print("Splitting references...")
        ref_para = paragraphs[refs_content_idx]

        # Collect all text from all runs
        full_ref_text = ''.join(run.text for run in ref_para.runs)

        # Split by newlines or by [N] pattern
        ref_items = re.split(r'\n+', full_ref_text)
        ref_items = [r.strip() for r in ref_items if r.strip()]

        # If splitting by newline didn't work well, split by [N] pattern
        if len(ref_items) <= 1:
            ref_items = re.findall(r'\[\d+\][^[]+', full_ref_text)
            ref_items = [r.strip() for r in ref_items if r.strip()]

        print(f"  Found {len(ref_items)} reference items")

        new_ref_paras = []
        for ref_text in ref_items:
            p = doc.add_paragraph()
            doc.element.body.remove(p._element)
            run = p.add_run(ref_text.strip())
            set_run_fonts(run, east_asian='宋体', size=SIZE_WU)
            set_paragraph_format(p, line_spacing=1.5)
            new_ref_paras.append(p)

        # Insert after references heading, remove old paragraph
        insert_target = paragraphs[refs_heading_idx]
        for p in reversed(new_ref_paras):
            insert_paragraph_after(insert_target, p)

        # Remove the original single-paragraph references
        ref_para._element.getparent().remove(ref_para._element)

    # ── 6. Insert funding and author bio after references ──
    print("Inserting funding and author bio...")
    # Find the last paragraph (should be a reference item after splitting)
    all_paras = list(doc.paragraphs)
    last_para = all_paras[-1] if all_paras else None

    end_paras = []
    # Funding
    p_fund = doc.add_paragraph()
    doc.element.body.remove(p_fund._element)
    run = p_fund.add_run(FUNDING_TEXT)
    set_run_fonts(run, east_asian='宋体', size=SIZE_WU)
    end_paras.append(p_fund)

    # Author bio
    p_bio = doc.add_paragraph()
    doc.element.body.remove(p_bio._element)
    run = p_bio.add_run(AUTHOR_BIO_TEXT)
    set_run_fonts(run, east_asian='宋体', size=SIZE_WU)
    end_paras.append(p_bio)

    if last_para is not None:
        for p in reversed(end_paras):
            insert_paragraph_after(last_para, p)

    # ── 7. Format tables ───────────────────────────────
    print(f"Formatting {len(doc.tables)} tables...")
    for table in doc.tables:
        # Table alignment
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.line_spacing = 1.0
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        set_run_fonts(run, east_asian='宋体', size=SIZE_WU)
                        if ri == 0:  # Header row bold
                            run.font.bold = True

        # Set table borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        # Remove existing borders
        existing = tblPr.findall(qn('w:tblBorders'))
        for e in existing:
            tblPr.remove(e)

        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        tblPr.append(borders)

    # ── 8. Apply superscript to in-text citations ──────
    print("Applying superscript to in-text citations...")
    all_paras = list(doc.paragraphs)
    # Full citation match with brackets (group capture for split to preserve brackets)
    citation_pattern = re.compile(r'\[\d+(?:[-–]\d+)?(?:[,，]\s*\d+(?:[-–]\d+)?)*\]')

    for p in all_paras:
        text = p.text.strip()
        # Skip non-body paragraphs by content pattern
        if not text:  # empty paragraph
            continue
        if text.startswith('关键') or text.startswith('關键'):
            continue
        if is_heading_1(text) or is_heading_2(text) or is_heading_3(text):
            continue
        if is_reference_item(text):
            continue
        if text == '参考文献' or text == '參考文献':
            continue
        if is_table_caption(text):
            continue
        if is_figure_caption(text):
            continue
        if text.startswith('中图分类号') or text.startswith('Abstract') or text.startswith('Key words'):
            continue
        if text.startswith('作者简介') or text.startswith('基金项目'):
            continue
        # English-only sections
        if text.startswith('Vision-Semantic') or text.startswith('AUTHOR'):
            continue

        # Process each run (use a copy of runs since we may modify the list)
        for run in list(p.runs):
            run_text = run.text
            if not citation_pattern.search(run_text):
                continue

            # Simple case: the entire run is a single citation
            if citation_pattern.fullmatch(run_text.strip()):
                run.font.superscript = True
                # Replace full-width hyphens in citation ranges
                run.text = run_text.replace('–', '-').replace('－', '-')
                continue

            # Complex case: citation embedded — split preserving brackets via group capture
            parts = re.split(r'(\[\d+(?:[-–]\d+)?(?:[,，]\s*\d+(?:[-–]\d+)?)*\])', run_text)
            if len(parts) <= 1:
                continue

            # Build new runs
            parent = run._element.getparent()
            run_index = list(parent).index(run._element)
            parent.remove(run._element)

            insert_pos = run_index
            for part in parts:
                if not part:
                    continue
                new_r = OxmlElement('w:r')
                orig_rPr = run._element.find(qn('w:rPr'))
                if orig_rPr is not None:
                    new_rPr = copy.deepcopy(orig_rPr)
                    new_r.append(new_rPr)

                new_t = OxmlElement('w:t')
                new_t.text = part
                new_t.set(qn('xml:space'), 'preserve')
                new_r.append(new_t)

                # If this part is a citation (starts with [ and matches citation pattern)
                if part.startswith('[') and citation_pattern.fullmatch(part):
                    rPr = new_r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        new_r.insert(0, rPr)
                    vertAlign = OxmlElement('w:vertAlign')
                    vertAlign.set(qn('w:val'), 'superscript')
                    rPr.append(vertAlign)
                    new_t.text = part.replace('–', '-').replace('－', '-')

                parent.insert(insert_pos, new_r)
                insert_pos += 1

    # ── 9. Save ─────────────────────────────────────────
    print(f"Saving to {DST}...")
    doc.save(DST)
    print("Done! Formatted document saved.")
    print(f"\nSummary:")
    print(f"  - Page: A4, margins 2.54cm, single column")
    print(f"  - Fonts: Chinese 宋体/黑体, Western Times New Roman")
    print(f"  - Body: 小四号(12pt) 宋体, 1.5x line spacing")
    print(f"  - Title: 二号(22pt) 黑体, centered")
    print(f"  - Headings: L1=四号(14pt)黑体, L2/L3=小四号(12pt)黑体")
    print(f"  - Tables: {len(doc.tables)} tables formatted")
    print(f"  - Placeholders inserted (marked [待补充]/[待确认])")
    print(f"  - References: split into individual items")
    print(f"  - Citations: superscript applied")


if __name__ == '__main__':
    main()
