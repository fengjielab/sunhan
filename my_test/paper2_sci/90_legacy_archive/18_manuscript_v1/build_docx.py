from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(32, 55, 72)
MUTED = RGBColor(96, 108, 118)
LIGHT = "F4F6F9"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[i]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])


def set_east_asia(run, font_name):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def plain_math(text):
    text = re.sub(r"\\operatorname\{clip\}\\\{(.*)\\\}", r"clip(\1)", text)
    def balanced_group(source, brace_pos):
        depth = 0
        for idx in range(brace_pos, len(source)):
            if source[idx] == "{":
                depth += 1
            elif source[idx] == "}":
                depth -= 1
                if depth == 0:
                    return source[brace_pos + 1 : idx], idx + 1
        return None, brace_pos + 1

    while "\\frac{" in text:
        start = text.find("\\frac{")
        numerator, after_num = balanced_group(text, start + len("\\frac"))
        if numerator is None or after_num >= len(text) or text[after_num] != "{":
            break
        denominator, after_den = balanced_group(text, after_num)
        if denominator is None:
            break
        text = text[:start] + f"({numerator})/({denominator})" + text[after_den:]
    while "\\sqrt{" in text:
        start = text.find("\\sqrt{")
        radicand, after = balanced_group(text, start + len("\\sqrt"))
        if radicand is None:
            break
        text = text[:start] + f"√({radicand})" + text[after:]
    text = text.replace("\\hat{w}", "ŵ")
    text = text.replace("\\left", "").replace("\\right", "")
    text = text.replace("\\operatorname{clip}", "clip")
    text = text.replace("\\max", "max").replace("\\sqrt", "√")
    text = text.replace("\\times", "×").replace("\\int", "∫")
    text = text.replace("\\mu", "μ").replace("\\sigma", "σ").replace("\\zeta", "ζ").replace("\\tau", "τ")
    text = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\text\{([^{}]+)\}", r"\1", text)
    text = text.replace("\\,", " ").replace("\\;", " ")
    text = text.replace("\\[", "").replace("\\]", "").replace("\\(", "").replace("\\)", "")
    text = text.replace("\\{", "{").replace("\\}", "}")
    text = re.sub(r"\\([A-Za-z]+)", r"\1", text)
    text = text.replace("{", "").replace("}", "")
    return text.strip()


INLINE = re.compile(r"(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+`|\\\(.*?\\\))")


def add_inline(paragraph, text, zh=False, size=None):
    pos = 0
    for match in INLINE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            if zh:
                set_east_asia(run, "Microsoft YaHei")
            if size:
                run.font.size = Pt(size)
        token = match.group(0)
        if token.startswith("**"):
            content = token[2:-2]
            run = paragraph.add_run(content)
            run.bold = True
        elif token.startswith("*"):
            content = token[1:-1]
            run = paragraph.add_run(content)
            run.italic = True
        elif token.startswith("`"):
            content = token[1:-1]
            run = paragraph.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5 if not size else size)
        else:
            content = plain_math(token)
            run = paragraph.add_run(content)
            run.font.name = "Cambria Math"
        if zh and not token.startswith("`"):
            set_east_asia(run, "Microsoft YaHei")
        if size:
            run.font.size = Pt(size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if zh:
            set_east_asia(run, "Microsoft YaHei")
        if size:
            run.font.size = Pt(size)


def configure_styles(doc, zh=False):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    if zh:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        if zh:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("Caption",):
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(9)
        st.font.italic = False
        st.font.color.rgb = NAVY
        if zh:
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.paragraph_format.space_before = Pt(3)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.keep_with_next = False

    if "Internal Note" not in styles:
        note = styles.add_style("Internal Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Internal Note"]
    note.font.name = "Calibri"
    note.font.size = Pt(9.5)
    note.font.color.rgb = MUTED
    if zh:
        note._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    note.paragraph_format.left_indent = Inches(0.18)
    note.paragraph_format.right_indent = Inches(0.18)
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(10)

    if "Equation" not in styles:
        eq = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        eq = styles["Equation"]
    eq.font.name = "Cambria Math"
    eq.font.size = Pt(10.5)
    eq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_before = Pt(4)
    eq.paragraph_format.space_after = Pt(6)
    eq.paragraph_format.keep_together = True


def add_note_shading(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2E74B5")
    pBdr.append(left)
    pPr.append(pBdr)


def build(md_path: Path, docx_path: Path, zh=False):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.different_first_page_header_footer = True
    configure_styles(doc, zh=zh)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("REALIZED-INTERVENTION AUDIT — MANUSCRIPT V1" if not zh else "实际干预审计——论文第一版")
    hr.font.name = "Calibri"
    hr.font.size = Pt(8.5)
    hr.font.color.rgb = MUTED
    if zh:
        set_east_asia(hr, "Microsoft YaHei")
    add_page_number(section.footer.paragraphs[0])

    in_equation = False
    equation_lines = []
    i = 0
    in_references = False
    title_done = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped == "\\[":
            in_equation = True
            equation_lines = []
            i += 1
            continue
        if in_equation:
            if stripped == "\\]":
                p = doc.add_paragraph(style="Equation")
                r = p.add_run(plain_math(" ".join(equation_lines)))
                r.font.name = "Cambria Math"
                r.font.size = Pt(10.5)
                in_equation = False
            else:
                equation_lines.append(stripped)
            i += 1
            continue

        if stripped.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
            if m:
                image_path = (md_path.parent / m.group(2)).resolve()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True
                inline_shape = p.add_run().add_picture(str(image_path), width=Inches(6.15))
                inline_shape._inline.docPr.set("descr", m.group(1))
                inline_shape._inline.docPr.set("title", m.group(1))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1].strip()):
            table_lines = [stripped]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in row.strip("|").split("|")] for row in table_lines]
            ncols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for r_idx, row in enumerate(rows):
                for c_idx, value in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.08
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 or r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    cell_font_size = 7.4 if (r_idx == 0 and c_idx == 0 and not zh) else 8.2
                    add_inline(p, value, zh=zh, size=cell_font_size)
                    if r_idx == 0:
                        shade_cell(cell, LIGHT)
                        for run in p.runs:
                            run.bold = True
            set_repeat_table_header(table.rows[0])
            if ncols == 5:
                widths = [1100, 1880, 1980, 2150, 2250]
            else:
                widths = [9360 // ncols] * ncols
                widths[-1] += 9360 - sum(widths)
            set_table_geometry(table, widths)
            after = doc.add_paragraph()
            after.paragraph_format.space_after = Pt(2)
            continue

        if stripped.startswith("# "):
            heading = stripped[2:].strip()
            if not title_done:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(22)
                p.paragraph_format.space_after = Pt(12)
                r = p.add_run(heading)
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(22)
                r.font.color.rgb = NAVY
                if zh:
                    set_east_asia(r, "Microsoft YaHei")
                title_done = True
            else:
                if heading in ("References", "参考文献"):
                    in_references = True
                p = doc.add_paragraph(heading, style="Heading 1")
            i += 1
            continue
        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            p = doc.add_paragraph(heading, style="Heading 2")
            i += 1
            continue
        if stripped.startswith("### "):
            heading = stripped[4:].strip()
            p = doc.add_paragraph(heading, style="Heading 3")
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Internal Note")
            add_note_shading(p)
            add_inline(p, stripped[2:], zh=zh)
            i += 1
            continue

        if in_references and re.match(r"^\d+\. ", stripped):
            content = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.12
            add_inline(p, content, zh=zh, size=9.3)
            i += 1
            continue

        p = doc.add_paragraph()
        if stripped.startswith("**Figure ") or stripped.startswith("**图") or stripped.startswith("**Table ") or stripped.startswith("**表"):
            p.style = doc.styles["Caption"]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
        add_inline(p, stripped, zh=zh)
        i += 1

    core = doc.core_properties
    core.title = lines[0].lstrip("# ")
    core.subject = "Integrated SCI manuscript draft — realized-intervention audit"
    core.author = "Authors to be inserted"
    core.keywords = ""
    doc.save(docx_path)


if __name__ == "__main__":
    build(OUT / "manuscript_v1_en.md", OUT / "manuscript_v1_en.docx", zh=False)
    build(OUT / "manuscript_v1_zh.md", OUT / "manuscript_v1_zh.docx", zh=True)
    print(OUT / "manuscript_v1_en.docx")
    print(OUT / "manuscript_v1_zh.docx")
